# =========================================================
# EduDOC – Compliance Engine (Step 13 FINAL)
# Checks:
#   ✓ Page count
#   ✓ Font family
#   ✓ Font size
#   ✓ Line spacing
#   ✓ Margins
#   ✓ Heading structure
# =========================================================

import os
from typing import Dict, Any, List
from PyPDF2 import PdfReader
import docx


# =========================================================
# UTIL — PAGE COUNT
# =========================================================
def get_page_count(path: str) -> int:

    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        reader = PdfReader(path)
        return len(reader.pages)

    if ext == ".docx":
        doc = docx.Document(path)
        words = sum(len(p.text.split()) for p in doc.paragraphs)
        return max(1, words // 300)

    return 0


# =========================================================
# MAIN ENTRY
# =========================================================
def run_compliance(document_path: str, rules: Dict) -> Dict[str, Any]:

    issues: List[str] = []
    score = 100

    if not rules:
        return {
            "score": 100,
            "issues": [],
            "message": "No rules configured"
        }

    ext = os.path.splitext(document_path)[1].lower()

    # =====================================================
    # 1️⃣ PAGE COUNT
    # =====================================================
    page_count = get_page_count(document_path)

    page_rules = rules.get("page", {})

    min_pages = page_rules.get("min_pages")
    max_pages = page_rules.get("max_pages")

    if min_pages is not None and page_count < min_pages:
        issues.append(f"Page count too low ({page_count}). Min {min_pages}")
        score -= 20

    if max_pages is not None and page_count > max_pages:
        issues.append(f"Page count too high ({page_count}). Max {max_pages}")
        score -= 20


    # =====================================================
    # DOCX ONLY CHECKS
    # =====================================================
    if ext == ".docx":

        doc = docx.Document(document_path)

        # =================================================
        # 2️⃣ TYPOGRAPHY (font + size + spacing)
        # =================================================
        typo = rules.get("typography", {})

        expected_font = typo.get("font_name")
        expected_size = typo.get("font_size")
        expected_spacing = typo.get("line_spacing")

        for p_index, para in enumerate(doc.paragraphs):

            # ---------- Line spacing ----------
            # ---------- Line spacing ----------
            if expected_spacing is not None:
                spacing = para.paragraph_format.line_spacing

                if spacing is not None:
                    if float(spacing) != float(expected_spacing):
                        issues.append(
                            f"Line spacing mismatch at paragraph {p_index+1}: {spacing} ≠ {expected_spacing}"
                        )
                        score -= 2

            # ---------- Font checks ----------
            for run in para.runs:

                if expected_font:
                    run_font = (run.font.name or para.style.font.name or "").strip()

                    if run_font and run_font != expected_font:
                        issues.append(
                            f"Font mismatch at paragraph {p_index+1}: {run_font} ≠ {expected_font}"
                        )
                        score -= 2
                        break

                if expected_size is not None and run.font.size and run.font.size.pt:
                    size_pt = run.font.size.pt

                    if int(size_pt) != int(expected_size):
                        issues.append(
                            f"Font size mismatch at paragraph {p_index+1}: {size_pt}pt ≠ {expected_size}pt"
                        )
                        score -= 2
                        break


        # =================================================
        # 3️⃣ MARGIN CHECK
        # =================================================
        margin_rules = rules.get("margin", {})
        expected_margin = margin_rules.get("margin_inches")

        if expected_margin is not None:


            for section in doc.sections:

                margins = [
                    section.top_margin.inches,
                    section.bottom_margin.inches,
                    section.left_margin.inches,
                    section.right_margin.inches,
                ]

                for m in margins:
                    if round(m, 1) != float(expected_margin):
                        issues.append(
                            f"Margin mismatch: {m}in ≠ {expected_margin}in"
                        )
                        score -= 5
                        break



        # =================================================
        # 4️⃣ HEADING STRUCTURE CHECK
        # =================================================
        structure = rules.get("structure", {})
        require_headings = structure.get("require_headings")

        if require_headings:

            found = any(
                para.style.name.lower().startswith("heading")
                for para in doc.paragraphs
            )

            if not found:
                issues.append("No headings detected (Heading styles missing)")
                score -= 10


    # =====================================================
    # FINALIZE
    # =====================================================
    score = max(score, 0)

    return {
        "score": score,
        "issues": issues,
        "page_count": page_count
    }
