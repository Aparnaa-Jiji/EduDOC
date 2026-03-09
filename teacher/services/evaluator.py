# =========================================================
# DOCUMENT EVALUATOR – EduDOC (Production Grade)
# Rule-driven • Deterministic • Safe • Pure Processing
# =========================================================

from docx import Document
import logging
from pathlib import Path
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from teacher.services.internet_search import InternetSearchService
from spellchecker import SpellChecker
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX


# =====================================================
# PROFESSIONAL REVIEWER COMMENT SYSTEM
# =====================================================

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class DocumentEvaluator:

    PARAGRAPHS_PER_PAGE = 30

    # =====================================================
    # INIT
    # =====================================================
    def __init__(self, file_path, rules=None, plagiarism_enabled=True):

        self.file_path = file_path
        self.rules = rules or {}
        self.plagiarism_enabled = plagiarism_enabled

        self.issues = []
        self.total_checks = 0
        self.failed_checks = 0

        self.doc = None

    # =====================================================
    # SAFE HELPER (Required by extended rules)
    # =====================================================
    def _add_issue(self, rule, message):
        self.total_checks += 1
        self.failed_checks += 1

        self.issues.append({
            "rule": rule,
            "expected": "Valid",
            "found": message,
            "status": "FAIL"
        })

    # =====================================================
    # MAIN ENTRY
    # =====================================================
    def evaluate(self):

        path = Path(self.file_path)

        if not path.exists():
            return self._empty_result("File not found")

        if path.suffix.lower() != ".docx":
            return self._empty_result("Unsupported format")

        try:
            self.doc = Document(self.file_path)
        except Exception:
            logger.exception("Failed to open document")
            return self._empty_result("Unable to read document")

        # Core checks
        self.check_page_count()
        self.check_margins()
        self.check_font_family()
        self.check_section_order()
        self.check_page_numbering()
        self.check_word_and_letter_count()
        self.check_reference_format()

        # Extended checks
        self.evaluate_header_footer()
      
        self._check_font_sizes()
        self._check_chapter_numbering()
        self._check_subheading_numbering()
        self._check_grammar()
        self._check_borders()

        # Spelling (INFO level)
        spelling_result = self.check_spelling()
        self.issues.append(spelling_result)

        compliance = self._calculate_compliance()

        plagiarism_percent = (
            self._plagiarism_stub()
            if self.plagiarism_enabled
            else 0
        )

        return {
            "compliance_percent": compliance,
            "plagiarism_percent": plagiarism_percent,
            "issues": self.issues,
        }

    # =====================================================
    # HELPERS
    # =====================================================
    def _empty_result(self, error=None):

        if error:
            self.issues.append({
                "rule": "Document",
                "status": "ERROR",
                "message": error
            })

        return {
            "compliance_percent": 0,
            "plagiarism_percent": 0,
            "issues": self.issues
        }

    def _calculate_compliance(self):

        if self.total_checks == 0:
            return 100

        passed = self.total_checks - self.failed_checks
        return int((passed / self.total_checks) * 100)

    def _record(self, rule, expected, found, passed):

        self.total_checks += 1

        if not passed:
            self.failed_checks += 1

        self.issues.append({
            "rule": rule,
            "expected": expected,
            "found": found,
            "status": "PASS" if passed else "FAIL"
        })

    # =====================================================
    # GRAMMAR CHECK
    # =====================================================
    def _check_grammar(self):

        rule = self.rules.get("grammar_rules", {})
        if not rule.get("enabled"):
            return

        grammar_issues = []

        for idx, paragraph in enumerate(self.doc.paragraphs):

            text = paragraph.text.strip()
            if not text:
                continue

            if not text[0].isupper():
                grammar_issues.append(f"P{idx+1}: no capital start")

            if text[-1] not in [".", "?", "!"]:
                grammar_issues.append(f"P{idx+1}: missing punctuation")

            if "  " in text:
                grammar_issues.append(f"P{idx+1}: double space")

            words = text.lower().split()
            for i in range(len(words) - 1):
                if words[i] == words[i + 1]:
                    grammar_issues.append(f"P{idx+1}: repeated '{words[i]}'")

            if re.search(r"\s+[.,!?]", text):
                grammar_issues.append(f"P{idx+1}: space before punctuation")

            if len(words) > 40:
                grammar_issues.append(f"P{idx+1}: sentence too long")

        if grammar_issues:
            self._record("Grammar", "Correct grammar",
                         f"{len(grammar_issues)} issues", False)

            self.issues.append({
                "rule": "Grammar Details",
                "status": "INFO",
                "details": grammar_issues
            })
        else:
            self._record("Grammar", "Correct grammar", "No issues", True)

    # =====================================================
    # SPELLING
    # =====================================================
    def extract_words_for_spellcheck(self):

        words = []

        if not self.file_path.endswith(".docx"):
            return words

        document = Document(self.file_path)

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            text = re.sub(r"[^\w\s]", "", text)
            tokens = text.split()

            for token in tokens:
                token = token.lower()
                if token.isalpha() and len(token) > 2:
                    words.append(token)

        return words

    def check_spelling(self):

        spell = SpellChecker()
        words = self.extract_words_for_spellcheck()

        if not words:
            return {
                "rule": "Spelling",
                "status": "PASS",
                "severity": "INFO",
                "misspelled_count": 0,
                "misspelled_words": []
            }

        filtered = [w for w in words if not any(c.isdigit() for c in w)]
        misspelled = spell.unknown(filtered)

        return {
            "rule": "Spelling",
            "status": "PASS" if not misspelled else "INFO",
            "severity": "INFO",
            "misspelled_count": len(misspelled),
            "misspelled_words": sorted(list(misspelled))[:50]
        }

    # =====================================================
    # PLAGIARISM
    # =====================================================
    def _plagiarism_stub(self):

        if not self.doc:
            return 0

        full_text = " ".join(
            p.text.strip()
            for p in self.doc.paragraphs
            if p.text.strip()
        )

        if not full_text:
            return 0

        sources = InternetSearchService.search_sources(full_text, max_results=3)
        if not sources:
            return 0

        student_sentences = set(
            s.strip().lower()
            for s in re.split(r'[.!?]', full_text)
            if len(s.strip()) > 20
        )

        matched_count = 0

        for source in sources:
            try:
                import requests
                response = requests.get(source["url"], timeout=5)
                page_text = response.text.lower()

                for sentence in student_sentences:
                    if sentence in page_text:
                        matched_count += 1
            except Exception:
                continue

        if not student_sentences:
            return 0

        similarity_percent = int(
            (matched_count / len(student_sentences)) * 100
        )

        return max(1, min(similarity_percent, 100))
    
    def _check_font_sizes(self):

        rules = self.rules.get("font_size_rules", {})
        heading_size = rules.get("heading")
        subheading_size = rules.get("subheading")
        content_size = rules.get("content")

        if not any([heading_size, subheading_size, content_size]):
            return

        heading_violations = 0
        subheading_violations = 0
        content_violations = 0

        for para in self.doc.paragraphs:

            style_name = para.style.name.lower() if para.style else ""

            for run in para.runs:

                if not run.font.size:
                    continue

                size = round(run.font.size.pt)

                # Heading 1
                if "heading 1" in style_name and heading_size:
                    if size != heading_size:
                        heading_violations += 1

                # Heading 2
                elif "heading 2" in style_name and subheading_size:
                    if size != subheading_size:
                        subheading_violations += 1

                # Content
                elif content_size:
                    if size != content_size:
                        content_violations += 1

        total_violations = heading_violations + subheading_violations + content_violations

        if total_violations > 0:
            self._record(
                "Font Size",
                f"H1:{heading_size}, H2:{subheading_size}, Content:{content_size}",
                f"{total_violations} violations",
                False
            )
        else:
            self._record(
                "Font Size",
                "Correct sizes",
                "No violations",
                True
            )

    def _check_chapter_numbering(self):

        rule = self.rules.get("chapter_numbering", {})
        if not rule.get("enabled"):
            return

        expected = rule.get("start_from", 1)
        detected_numbers = []

        for para in self.doc.paragraphs:

            match = re.match(r"^(\d+)\.", para.text.strip())
            if match:
                detected_numbers.append(int(match.group(1)))

        passed = True

        for number in detected_numbers:
            if number != expected:
                passed = False
                break
            expected += 1

        if not detected_numbers:
            passed = False

        self._record(
            "Chapter Numbering",
            "Sequential 1,2,3...",
            detected_numbers if detected_numbers else "None",
            passed
        )
    
    def _check_subheading_numbering(self):

        rule = self.rules.get("subheading_numbering", {})
        if not rule.get("enabled"):
            return

        pattern = r"^(\d+)\.(\d+)"
        structure = {}

        passed = True

        for para in self.doc.paragraphs:

            match = re.match(pattern, para.text.strip())
            if not match:
                continue

            chapter = int(match.group(1))
            section = int(match.group(2))

            if chapter not in structure:
                structure[chapter] = 1

            if section != structure[chapter]:
                passed = False
                break

            structure[chapter] += 1

        if not structure:
            passed = False

        self._record(
            "Subheading Numbering",
            "1.1, 1.2, 2.1 sequential",
            structure if structure else "None",
            passed
        )
    
    # =====================================================
    # PAGE COUNT
    # =====================================================
    def check_page_count(self):

        rules = self.rules.get("page_rules", {})

        min_pages = rules.get("min_pages")
        max_pages = rules.get("max_pages")

        if not min_pages and not max_pages:
            return

        approx_pages = max(
            1,
            round(len(self.doc.paragraphs) / self.PARAGRAPHS_PER_PAGE)
        )

        passed = True

        if min_pages and approx_pages < min_pages:
            passed = False

        if max_pages and approx_pages > max_pages:
            passed = False

        self._record(
            "Page Count",
            f"{min_pages}-{max_pages} pages",
            f"{approx_pages} pages",
            passed
        )

    def check_margins(self):

        margins = self.rules.get("margins", {})
        if not margins:
            return

        section = self.doc.sections[0]

        found = {
            "top": round(section.top_margin.inches,2),
            "bottom": round(section.bottom_margin.inches,2),
            "left": round(section.left_margin.inches,2),
            "right": round(section.right_margin.inches,2),
        }

        passed = (
            abs(found["top"]-margins.get("top",0)) < 0.2 and
            abs(found["bottom"]-margins.get("bottom",0)) < 0.2 and
            abs(found["left"]-margins.get("left",0)) < 0.2 and
            abs(found["right"]-margins.get("right",0)) < 0.2
        )

        self._record("Margins", margins, found, passed)

    def check_font_family(self):

        rule = self.rules.get("font_rules", {})
        expected = rule.get("font_name")

        if not expected:
            return

        violations = 0

        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.name and run.font.name != expected:
                    violations += 1

        self._record(
            "Font Family",
            expected,
            f"{violations} violations",
            violations == 0
        )

    def check_section_order(self):

        expected = self.rules.get("sections", [])
        if not expected:
            return

        text = "\n".join(p.text.lower() for p in self.doc.paragraphs)

        missing = [
            sec for sec in expected
            if sec.lower() not in text
        ]

        self._record(
            "Section Order",
            expected,
            f"Missing: {missing}" if missing else "All found",
            len(missing)==0
        )

    def check_page_numbering(self):

        rule = self.rules.get("page_number", {})
        if not rule.get("enabled"):
            return

        footer_text = ""

        for section in self.doc.sections:
            footer_text += section.footer.paragraphs[0].text

        detected = bool(re.search(r"\d+", footer_text))

        self._record(
            "Page Number",
            "Required",
            "Detected" if detected else "Not Found",
            detected
        )

    def check_word_and_letter_count(self):

        text = " ".join(p.text for p in self.doc.paragraphs)

        words = len(text.split())
        letters = len(re.sub(r"\s","",text))

        self.issues.append({
            "rule":"Word & Letter Count",
            "status":"INFO",
            "found":{
                "word_count":words,
                "letter_count":letters
            }
        })

    def check_reference_format(self):

        text = "\n".join(p.text.lower() for p in self.doc.paragraphs)

        has_reference = "reference" in text

        self._record(
            "References Section",
            "Present",
            "Found" if has_reference else "Missing",
            has_reference
        )

    def evaluate_header_footer(self):

        rules = self.rules.get("layout_rules", {})

        header_rule = rules.get("header", {})
        footer_rule = rules.get("footer", {})

        header_text = ""
        footer_text = ""

        for sec in self.doc.sections:
            if sec.header.paragraphs:
                header_text += sec.header.paragraphs[0].text

            if sec.footer.paragraphs:
                footer_text += sec.footer.paragraphs[0].text

        if header_rule.get("required"):
            self._record(
                "Header",
                header_rule.get("text"),
                header_text or "Not Found",
                header_rule.get("text","").lower() in header_text.lower()
            )

        if footer_rule.get("required"):
            self._record(
                "Footer",
                footer_rule.get("text"),
                footer_text or "Not Found",
                footer_rule.get("text","").lower() in footer_text.lower()
            )
    
    def _check_borders(self):

        rule = self.rules.get("layout_rules",{}).get("border_required")

        if not rule:
            return

        has_border=False

        for sec in self.doc.sections:
            if sec._sectPr.xpath("./w:pgBorders"):
                has_border=True

        self._record(
            "Page Border",
            "Required",
            "Detected" if has_border else "Missing",
            has_border
        )

    # =====================================================
    # ANNOTATED FILE GENERATION
    # =====================================================
    # =====================================================
# PROFESSIONAL ANNOTATED DOCUMENT GENERATOR
# =====================================================



    def generate_annotated_file(self, output_path):

        if not self.doc:
            self.doc = Document(self.file_path)

        doc = self.doc

        # =================================================
        # 1️⃣ INSERT COVER REVIEW SUMMARY
        # =================================================
        summary = doc.add_paragraph()
        summary.add_run("EduDOC Academic Compliance Review\n").bold = True
        summary.add_run(f"\nCompliance Score : {self._calculate_compliance()}%\n")
        summary.add_run(f"Plagiarism Index : {self._plagiarism_stub()}%\n")
        summary.add_run(f"Total Violations : {self.failed_checks}\n\n")

        summary.add_run("Major Issues:\n").bold = True

        for issue in self.issues:
            if issue.get("status") == "FAIL":
                summary.add_run(f"• {issue['rule']}\n")

        doc.add_page_break()

        # =================================================
        # 2️⃣ HELPER → ADD REVIEWER COMMENT BELOW PARAGRAPH
        # =================================================
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph


        def add_reviewer_note(paragraph, message):

            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)

            new_para = Paragraph(new_p, paragraph._parent)

            run = new_para.add_run(f"Reviewer Note: {message}")
            run.italic = True
            run.font.size = Pt(9)
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
        # =================================================
        # 3️⃣ APPLY STRUCTURED REVIEW COMMENTS
        # =================================================
        for issue in self.issues:

            rule = issue.get("rule")
            status = issue.get("status")

            if status != "FAIL":
                continue

            # -------------------------
            # Header Issues
            # -------------------------
            if rule == "Header":
                for section in doc.sections:
                    for para in section.header.paragraphs:
                        for run in para.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.PINK
                    if section.header.paragraphs:
                        add_reviewer_note(
                            section.header.paragraphs[0],
                            "Header does not match the required format or text."
                        )

            # -------------------------
            # Footer Issues
            # -------------------------
            elif rule == "Footer":
                for section in doc.sections:
                    for para in section.footer.paragraphs:
                        for run in para.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.PINK
                    if section.footer.paragraphs:
                        add_reviewer_note(
                            section.footer.paragraphs[0],
                            "Footer format or required content missing."
                        )

            # -------------------------
            # Font Size Violations
            # -------------------------
            elif rule == "Font Size":
                for para in doc.paragraphs:
                    if "heading" in para.style.name.lower():
                        for run in para.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        add_reviewer_note(
                            para,
                            "Heading font size does not comply with prescribed guidelines."
                        )

            # -------------------------
            # Chapter Numbering
            # -------------------------
            elif rule == "Chapter Numbering":
                for para in doc.paragraphs:
                    if para.text.strip().startswith("1.") or para.text.strip().startswith("2."):
                        for run in para.runs:
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        add_reviewer_note(
                            para,
                            "Chapter numbering must follow sequential structure starting from configured value."
                        )

            # -------------------------
            # Grammar Issues
            # -------------------------
            elif rule == "Grammar":
                for para in doc.paragraphs:
                    for run in para.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                    add_reviewer_note(
                        para,
                        "Grammar inconsistencies detected. Review sentence structure and punctuation."
                    )
                    break  # avoid flooding document

            # -------------------------
            # Margins
            # -------------------------
            elif rule == "Margins":
                first_para = doc.paragraphs[0]
                add_reviewer_note(
                    first_para,
                    "Page margins do not meet the required specifications."
                )

        # =================================================
        # 4️⃣ SPELLING (INFO LEVEL COMMENT)
        # =================================================
        for issue in self.issues:

            if issue.get("rule") == "Spelling":
                misspelled = issue.get("misspelled_words", [])

                if misspelled:
                    para = doc.paragraphs[0]
                    add_reviewer_note(
                        para,
                        f"Spelling issues detected: {', '.join(misspelled[:5])}."
                    )

        # =================================================
        # 5️⃣ SAVE FILE
        # =================================================
        doc.save(output_path)